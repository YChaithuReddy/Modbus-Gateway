"""
Convert SPOKEN_SCRIPT.md → SPOKEN_SCRIPT.docx (Word document).
Handles: headings, bold/italic, blockquotes, bullet lists, checkboxes, tables,
         horizontal rules, emoji in headings.
Optimized for print / offline rehearsal use.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- CONFIG ----------
HERE = Path(__file__).parent
SRC = HERE / "SPOKEN_SCRIPT.md"
DST = HERE / "SPOKEN_SCRIPT.docx"

# Colors (matching deck palette)
NAVY       = RGBColor(0x0A, 0x25, 0x40)
TEAL       = RGBColor(0x00, 0x8C, 0x85)
AMBER_DARK = RGBColor(0xD4, 0x84, 0x41)
TEXT_DARK  = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_MID   = RGBColor(0x5B, 0x64, 0x72)
DIVIDER    = RGBColor(0xE5, 0xE8, 0xEC)
QUOTE_BG   = RGBColor(0xF3, 0xF9, 0xF8)  # very subtle teal-tint


# ---------- HELPERS ----------
def set_cell_border(cell, color="E5E8EC", sz="4"):
    """Add cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def shade_cell(cell, color_hex):
    """Background-shade a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_paragraph_border_left(paragraph, color_hex="00C2B8", size="18"):
    """Add a left border (for blockquote styling)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def add_paragraph_shading(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


# ---------- INLINE FORMATTING PARSER ----------
INLINE_PATTERN = re.compile(r'(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)')


def add_runs(paragraph, text, *, base_bold=False, base_italic=False,
             base_color=None, base_font=None, base_size=None):
    """Add inline-formatted runs to a paragraph."""
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            if base_size:
                run.font.size = Pt(base_size.pt - 1 if hasattr(base_size, 'pt') else 10)
            else:
                run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)

        # Apply base formatting that wasn't overridden by inline
        if base_bold and not (part.startswith("**") and part.endswith("**")):
            run.bold = True
        if base_italic and not (part.startswith("*") and part.endswith("*") and not part.startswith("**")):
            run.italic = True
        if base_color:
            run.font.color.rgb = base_color
        if base_font and not (part.startswith("`") and part.endswith("`")):
            run.font.name = base_font
        if base_size and not (part.startswith("`") and part.endswith("`")):
            run.font.size = base_size


# ---------- DOCUMENT SETUP ----------
doc = Document()

# Tighten default margins
for section in doc.sections:
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

# Set Normal style base font
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = TEXT_DARK
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.25


# ---------- PARSER ----------
def read_md(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.readlines()


def is_table_separator(line):
    """Matches lines like |---|---|---|"""
    s = line.strip()
    if not s.startswith("|"):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    return all(re.fullmatch(r":?-+:?", c) for c in cells if c)


def parse_table_row(line):
    s = line.strip()
    if not s.startswith("|"):
        return None
    return [c.strip() for c in s.strip("|").split("|")]


def add_heading(doc, text, level):
    """Custom heading — python-docx default heading styles get overridden."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True

    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = TEAL
    return p


def add_blockquote(doc, lines):
    """Blockquote: indented, subtle teal background, left border, larger italic-free type."""
    text = "\n".join(lines).strip()
    if not text:
        return
    # Split on blank-line-ish separations
    # Markdown blockquotes in the script use `>` prefix, empty `>` lines are paragraph breaks
    chunks = re.split(r"\n\s*\n", text)
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.35
        add_paragraph_border_left(p, color_hex="00C2B8", size="24")
        add_paragraph_shading(p, "F3F9F8")
        # Remove > markers from line beginnings
        clean = re.sub(r"^>\s?", "", chunk, flags=re.MULTILINE)
        add_runs(p, clean, base_size=Pt(12))


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # Border-bottom on an empty paragraph creates a horizontal rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E8EC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullet(doc, text, *, checkbox=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(2)
    prefix_run = p.add_run(("☐  " if checkbox else "•  "))
    prefix_run.font.color.rgb = TEAL
    prefix_run.font.bold = True
    add_runs(p, text)


def add_table_block(doc, rows):
    """Create a Word table from parsed markdown rows."""
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            set_cell_border(cell)
            if r_idx == 0:
                shade_cell(cell, "0A2540")
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(cell_text)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
            else:
                if r_idx % 2 == 0:
                    shade_cell(cell, "F8F9FB")
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                add_runs(p, cell_text, base_size=Pt(10))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ---------- MAIN PARSE LOOP ----------
lines = read_md(SRC)
i = 0
in_quote = False
quote_lines = []

while i < len(lines):
    raw = lines[i]
    stripped = raw.rstrip("\n")

    # Blockquote accumulation
    if stripped.startswith(">"):
        quote_lines.append(stripped)
        i += 1
        continue
    else:
        if quote_lines:
            add_blockquote(doc, quote_lines)
            quote_lines = []

    # Empty line
    if not stripped.strip():
        i += 1
        continue

    # Horizontal rule
    if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", stripped.strip()):
        add_hr(doc)
        i += 1
        continue

    # Headings
    m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
    if m:
        level = len(m.group(1))
        text = m.group(2)
        add_heading(doc, text, min(level, 3))
        i += 1
        continue

    # Table
    if stripped.lstrip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
        rows = []
        header = parse_table_row(stripped)
        if header:
            rows.append(header)
        i += 2  # skip header + separator
        while i < len(lines) and lines[i].lstrip().startswith("|"):
            r = parse_table_row(lines[i])
            if r:
                rows.append(r)
            i += 1
        add_table_block(doc, rows)
        continue

    # Checkbox list item
    cb = re.match(r"^\s*[-*]\s+\[( |x|X)\]\s+(.+)$", stripped)
    if cb:
        add_bullet(doc, cb.group(2), checkbox=True)
        i += 1
        continue

    # Bullet list item
    bl = re.match(r"^\s*[-*]\s+(.+)$", stripped)
    if bl:
        add_bullet(doc, bl.group(1))
        i += 1
        continue

    # Numbered list
    nl = re.match(r"^\s*(\d+)\.\s+(.+)$", stripped)
    if nl:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.space_after = Pt(2)
        num_run = p.add_run(f"{nl.group(1)}.  ")
        num_run.font.bold = True
        num_run.font.color.rgb = TEAL
        add_runs(p, nl.group(2))
        i += 1
        continue

    # Regular paragraph
    p = doc.add_paragraph()
    add_runs(p, stripped)
    i += 1

# Flush any trailing blockquote
if quote_lines:
    add_blockquote(doc, quote_lines)


# ---------- SAVE ----------
doc.save(DST)
print(f"Generated: {DST}")
print(f"Size: {DST.stat().st_size // 1024} KB")
