"""Convert FULL_PROJECT_REPORT.md to a formatted Word document."""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE = os.path.join(SCRIPT_DIR, "FULL_PROJECT_REPORT.md")
DOCX_FILE = os.path.join(SCRIPT_DIR, "FluxGen_IoT_Gateway_Full_Project_Report.docx")


def set_cell_shading(cell, color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_document(doc):
    """Set up document styles."""
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title style
    style = doc.styles['Title']
    style.font.size = Pt(26)
    style.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    style.font.bold = True
    style.paragraph_format.space_after = Pt(6)

    # Heading 1
    style = doc.styles['Heading 1']
    style.font.size = Pt(18)
    style.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(24)
    style.paragraph_format.space_after = Pt(8)

    # Heading 2
    style = doc.styles['Heading 2']
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0x2d, 0x3a, 0x4a)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after = Pt(6)

    # Heading 3
    style = doc.styles['Heading 3']
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(4)

    # Normal text
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = 'Calibri'
    style.paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header.strip()
        set_cell_shading(cell, "1a56db")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = "f0f4ff" if r_idx % 2 == 0 else "ffffff"
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text.strip()
            set_cell_shading(cell, bg)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def add_code_block(doc, code_text):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2d, 0x3a, 0x4a)


def parse_and_convert(md_path, docx_path):
    """Parse markdown and create Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    style_document(doc)

    in_code_block = False
    code_buffer = []
    in_table = False
    table_headers = []
    table_rows = []
    skip_toc = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table and table_headers:
                    add_table(doc, table_headers, table_rows)
                    table_headers = []
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]

            # Check if separator row
            if all(re.match(r'^[-:]+$', c.strip()) for c in cells if c.strip()):
                i += 1
                continue

            if not in_table:
                # This is the header row
                in_table = True
                table_headers = cells
                table_rows = []
            else:
                table_rows.append(cells)

            i += 1
            continue
        else:
            # Flush pending table
            if in_table and table_headers:
                add_table(doc, table_headers, table_rows)
                table_headers = []
                table_rows = []
                in_table = False

        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip horizontal rules
        if stripped == '---':
            # Add a thin line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # TOC section - skip
        if stripped == '## Table of Contents':
            skip_toc = True
            i += 1
            continue
        if skip_toc:
            if stripped.startswith('## ') or stripped.startswith('# '):
                if not stripped.startswith('## Table'):
                    skip_toc = False
                else:
                    i += 1
                    continue
            else:
                i += 1
                continue

        # Title (# at top)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            # First title as document title
            if text == 'FluxGen Modbus IoT Gateway':
                p = doc.add_paragraph(text, style='Title')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif text == 'Complete Project Report':
                p = doc.add_paragraph(text, style='Subtitle')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, level=1)
            i += 1
            continue

        # Heading 2
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        # Heading 3
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # Heading 4
        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Bold metadata lines (like **Company**: FluxGen)
        if stripped.startswith('**') and '**:' in stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Parse bold:value pattern
            match = re.match(r'\*\*(.+?)\*\*:\s*(.*)', stripped)
            if match:
                run = p.add_run(match.group(1) + ': ')
                run.bold = True
                run.font.size = Pt(10)
                run2 = p.add_run(match.group(2))
                run2.font.size = Pt(10)
            else:
                p.add_run(stripped.replace('**', ''))
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold within bullets
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
            i += 1
            continue

        # Numbered items
        match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph - handle bold/italic
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10.5)
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.italic = True
                run.font.size = Pt(10.5)
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
            else:
                run = p.add_run(part)
                run.font.size = Pt(10.5)

        i += 1

    # Flush any remaining table
    if in_table and table_headers:
        add_table(doc, table_headers, table_rows)

    # Flush any remaining code block
    if code_buffer:
        add_code_block(doc, '\n'.join(code_buffer))

    # Add footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FluxGen Technologies | Industrial IoT Solutions')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    doc.save(docx_path)
    print(f"Word document saved: {docx_path}")
    print(f"Size: {os.path.getsize(docx_path) / 1024:.0f} KB")


if __name__ == "__main__":
    parse_and_convert(MD_FILE, DOCX_FILE)
